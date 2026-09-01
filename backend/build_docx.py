#!/usr/bin/env python3
"""
Word-dokumentin generointi.
Vastaanottaa markdown-muotoisen dokumenttitekstin ja luo tyylitellyn .docx-tiedoston.
Käyttö: python build_docx.py --stdin output.docx
        python build_docx.py --file input.json output.docx
"""
import sys
import json
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ═══ VÄRIT ═══
DEEP_BLUE = RGBColor(0x0C, 0x23, 0x40)
DIGITAL_BLUE = RGBColor(0x1B, 0x6C, 0xA8)
ORANGE = RGBColor(0xE8, 0x52, 0x1A)
MINT = RGBColor(0x3B, 0xBF, 0xAD)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x8C, 0x9B, 0xAA)
LIGHT = RGBColor(0xEE, 0xF1, 0xF3)


def set_cell_shading(cell, color_hex):
    """Aseta solun taustaväri."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elm)


def create_docx(data, out_path):
    """Luo Word-dokumentti markdown-tekstistä."""
    document_text = data.get("documentText", "")
    chapters = data.get("chapters", [])
    lang = data.get("lang", "fi")

    doc = Document()

    # ═══ TYYLIASETUKSET ═══
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DEEP_BLUE

    # Otsikkotyylit
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_font = h_style.font
        h_font.name = 'Calibri'
        h_font.color.rgb = DEEP_BLUE if level <= 2 else DIGITAL_BLUE
        h_font.size = Pt(24 - (level * 4))  # 20, 16, 12
        h_font.bold = True

    # Sivun asetukset
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ═══ KANSILEHTI ═══
    add_cover_page(doc, document_text, chapters)

    # ═══ SISÄLTÖ ═══
    if document_text:
        parse_markdown_to_docx(doc, document_text)
    else:
        # Fallback: ei tekstiä
        doc.add_paragraph("Dokumentti on tyhjä." if lang == "fi" else "Document is empty.")

    # ═══ LOPPUSIVU ═══
    doc.add_page_break()
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Pioneering\nan ethical\ndigital world.")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DEEP_BLUE

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = ORANGE
    run.font.letter_spacing = Pt(3)

    # Tallenna
    doc.save(out_path)
    print(f"OK: {out_path}")


def add_cover_page(doc, text, chapters):
    """Luo kansilehti dokumentin ensimmäisestä otsikosta tai chapters-datasta."""
    # Etsi otsikko tekstistä (ensimmäinen # -rivi)
    title = ""
    subtitle = ""
    first_h1 = re.search(r'^#\s+(.+)', text, re.MULTILINE)
    if first_h1:
        title = first_h1.group(1).strip()
    elif chapters:
        title_ch = next((c for c in chapters if c.get("layout") == "title"), None)
        title = title_ch.get("label", "Dokumentti") if title_ch else "Dokumentti"

    # Etsi alaotsikko (teksti heti otsikon jälkeen, ennen seuraavaa otsikkoa)
    if first_h1:
        after_title = text[first_h1.end():].strip()
        lines = after_title.split("\n")
        for line in lines:
            stripped = line.strip()
            if stripped and not stripped.startswith("#"):
                subtitle = stripped
                break

    # Oranssi yläviiva
    p = doc.add_paragraph()
    run = p.add_run("━" * 60)
    run.font.color.rgb = ORANGE
    run.font.size = Pt(6)

    for _ in range(4):
        doc.add_paragraph()

    # Otsikko
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title or "Dokumentti")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DEEP_BLUE

    # Alaotsikko
    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = ORANGE

    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = ORANGE
    run.font.letter_spacing = Pt(3)

    doc.add_page_break()


def parse_markdown_to_docx(doc, text):
    """Parsii markdown-tekstin ja lisää sen Word-dokumenttiin."""
    lines = text.split("\n")
    i = 0
    # Ohita ensimmäinen # -otsikko ja sitä seuraava alaotsikkorivi (kansilehti hoiti ne)
    first_h1_skipped = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Tyhjä rivi
        if not stripped:
            i += 1
            continue

        # H1 otsikko: # Otsikko
        if re.match(r'^#\s+', stripped):
            heading_text = re.sub(r'^#\s+', '', stripped).strip()
            # Ohita ensimmäinen H1 (kansilehti)
            if not first_h1_skipped:
                first_h1_skipped = True
                i += 1
                # Ohita myös seuraava ei-tyhjä rivi jos se on alaotsikko (ei # eikä -)
                while i < len(lines) and not lines[i].strip():
                    i += 1
                if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#"):
                    i += 1  # skip subtitle line
                continue

            heading_text = clean_markdown(heading_text)
            doc.add_heading(heading_text, level=1)
            i += 1
            continue

        # H2 otsikko: ## Otsikko
        if re.match(r'^##\s+', stripped):
            heading_text = re.sub(r'^##\s+', '', stripped).strip()
            heading_text = clean_markdown(heading_text)
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        # H3 otsikko: ### Otsikko
        if re.match(r'^###\s+', stripped):
            heading_text = re.sub(r'^###\s+', '', stripped).strip()
            heading_text = clean_markdown(heading_text)
            doc.add_heading(heading_text, level=3)
            i += 1
            continue

        # Taulukko: | sarake1 | sarake2 |
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            table_data = parse_table_lines(table_lines)
            if table_data and len(table_data) >= 2:
                add_styled_table(doc, table_data)
            continue

        # Bullet-lista: - item tai • item tai * item
        if re.match(r'^[-•*]\s+', stripped):
            while i < len(lines):
                l = lines[i].strip()
                if re.match(r'^[-•*]\s+', l):
                    item_text = re.sub(r'^[-•*]\s+', '', l).strip()
                    item_text = clean_markdown(item_text)
                    p = doc.add_paragraph(item_text, style='List Bullet')
                    p.paragraph_format.space_after = Pt(4)
                    i += 1
                elif not l:
                    i += 1
                    break
                else:
                    break
            continue

        # Numeroidut listat: 1. item
        if re.match(r'^\d+\.\s+', stripped):
            while i < len(lines):
                l = lines[i].strip()
                if re.match(r'^\d+\.\s+', l):
                    item_text = re.sub(r'^\d+\.\s+', '', l).strip()
                    item_text = clean_markdown(item_text)
                    p = doc.add_paragraph(item_text, style='List Number')
                    p.paragraph_format.space_after = Pt(4)
                    i += 1
                elif not l:
                    i += 1
                    break
                else:
                    break
            continue

        # Horisontaalinen viiva: --- tai ***
        if re.match(r'^[-*_]{3,}$', stripped):
            p = doc.add_paragraph()
            run = p.add_run("━" * 40)
            run.font.color.rgb = ORANGE
            run.font.size = Pt(6)
            i += 1
            continue

        # Normaali kappale — kerää rivit kunnes tyhjä rivi tai uusi elementti
        para_lines = []
        while i < len(lines):
            l = lines[i].strip()
            if not l:
                i += 1
                break
            if l.startswith("#") or l.startswith("|") or re.match(r'^[-•*]\s+', l) or re.match(r'^\d+\.\s+', l) or re.match(r'^[-*_]{3,}$', l):
                break
            para_lines.append(l)
            i += 1

        if para_lines:
            para_text = " ".join(para_lines)
            add_rich_paragraph(doc, para_text)


def clean_markdown(text):
    """Poista markdown-muotoilu tekstistä."""
    # Poista bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Poista italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Poista linkit
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    # Poista kooditagi
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()


def add_rich_paragraph(doc, text):
    """Lisää kappale jossa bold-teksti on lihavoitu."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)

    # Pilko bold-osiin: **teksti** → lihavoitu
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = DEEP_BLUE
        else:
            # Poista jäljellä olevat markdown-merkinnät
            cleaned = clean_markdown(part)
            if cleaned:
                run = p.add_run(cleaned)
                run.font.color.rgb = DEEP_BLUE


def parse_table_lines(lines):
    """Parsii markdown-taulukkorivit dataksi."""
    # Poista separator-rivit (---|---)
    data_lines = [l for l in lines if not re.match(r'^\|[\s\-:|]+\|$', l)]
    if len(data_lines) < 2:
        return None
    result = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        cells = [clean_markdown(c) for c in cells]
        result.append(cells)
    return result


def add_styled_table(doc, table_data):
    """Luo tyylitelty taulukko """
    if not table_data or len(table_data) < 2:
        return

    num_cols = len(table_data[0])
    table = doc.add_table(rows=len(table_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Otsikkorivi
    for ci, cell_text in enumerate(table_data[0]):
        cell = table.rows[0].cells[ci]
        cell.text = str(cell_text).strip()
        set_cell_shading(cell, "0C2340")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)

    # Datarivit
    for ri in range(1, len(table_data)):
        bg = "FFFFFF" if ri % 2 == 1 else "EEF1F3"
        for ci in range(min(len(table_data[ri]), num_cols)):
            cell = table.rows[ri].cells[ci]
            cell.text = str(table_data[ri][ci]).strip()
            set_cell_shading(cell, bg)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = DEEP_BLUE

    doc.add_paragraph()


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--stdin", action="store_true")
    parser.add_argument("--file", type=str, default=None)
    parser.add_argument("output", type=str)
    args = parser.parse_args()

    if args.file:
        with open(args.file, "r", encoding="utf-8") as f:
            data = json.load(f)
    elif args.stdin:
        data = json.load(sys.stdin)
    else:
        print("Usage: --stdin or --file <path>", file=sys.stderr)
        sys.exit(1)

    create_docx(data, args.output)
